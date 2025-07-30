import streamlit as st
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
import pytesseract
from googletrans import Translator
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import pandas as pd
import io
import openai
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
import os
from deep_translator import GoogleTranslator
import sys
import locale
import fitz  # PyMuPDF para leer PDFs
import unicodedata
import re
import requests
import urllib.request
from pathlib import Path
import cv2
import numpy as np

# Configura la página
st.set_page_config(page_title="Translator de Documentos", layout="wide")

openai_api = st.text_input("OpenAI API Key", key="chatbot_api_key", type="password")
if not openai_api:
    st.info("Por favor escriba su OpenAI API key para continuar.")
    st.stop()

OPENAI_API_KEY = openai_api

st.title("🈳 OCR y Traducción Multiidioma desde Imágenes y PDFs")
st.write("Sube una imagen o PDF con texto en cualquier idioma. La app extraerá el texto y lo traducirá al inglés.")

# Diccionario de idiomas soportados con sus códigos ISO
IDIOMAS_TESSERACT = {
    'Árabe': 'ara',
    'Chino Simplificado': 'chi_sim',
    'Chino Tradicional': 'chi_tra',
    'Hebreo': 'heb',
    'Japonés': 'jpn',
    'Coreano': 'kor',
    'Ruso': 'rus',
    'Hindi': 'hin',
    'Tailandés': 'tha',
    'Vietnamita': 'vie',
    'Inglés': 'eng',
    'Español': 'spa',
    'Francés': 'fra',
    'Alemán': 'deu',
    'Italiano': 'ita',
    'Portugués': 'por',
    'Detección Automática': 'auto'
}


def verificar_idiomas_instalados():
    """
    Verifica qué idiomas están instalados en Tesseract
    """
    try:
        idiomas_disponibles = pytesseract.get_languages(config='')
        return idiomas_disponibles
    except Exception as e:
        st.error(f"Error al verificar idiomas: {str(e)}")
        return ['eng']  # Solo inglés por defecto


def descargar_paquete_idioma(codigo_idioma):
    """
    Descarga automáticamente un paquete de idioma para Tesseract
    """
    try:
        # URL base para descargar paquetes de idioma
        base_url = "https://github.com/tesseract-ocr/tessdata/raw/main/"
        filename = f"{codigo_idioma}.traineddata"

        # Intentar encontrar la carpeta tessdata
        posibles_rutas = [
            r"C:\Program Files\Tesseract-OCR\tessdata",
            r"C:\Program Files (x86)\Tesseract-OCR\tessdata",
            "/usr/share/tesseract-ocr/4.00/tessdata",
            "/usr/share/tesseract-ocr/tessdata",
            "/opt/homebrew/share/tessdata"
        ]

        tessdata_path = None
        for ruta in posibles_rutas:
            if os.path.exists(ruta):
                tessdata_path = ruta
                break

        if not tessdata_path:
            st.error("No se pudo encontrar la carpeta tessdata de Tesseract")
            return False

        archivo_destino = os.path.join(tessdata_path, filename)

        if os.path.exists(archivo_destino):
            return True  # Ya está instalado

        # Descargar el archivo
        url = base_url + filename
        with st.spinner(f"Descargando paquete de idioma {codigo_idioma}..."):
            urllib.request.urlretrieve(url, archivo_destino)

        st.success(f"Paquete de idioma {codigo_idioma} descargado exitosamente")
        return True

    except Exception as e:
        st.error(f"Error al descargar paquete de idioma {codigo_idioma}: {str(e)}")
        return False


def limpiar_texto_para_pdf(texto):
    """
    Limpia el texto para evitar problemas de codificación en PDF
    """
    if not texto:
        return ""

    # Normalizar caracteres Unicode
    texto = unicodedata.normalize('NFKD', texto)

    # Reemplazar caracteres problemáticos comunes
    replacements = {
        '\u201c': '"',  # comilla izquierda
        '\u201d': '"',  # comilla derecha
        '\u2018': "'",  # comilla simple izquierda
        '\u2019': "'",  # comilla simple derecha
        '\u2013': '-',  # en dash
        '\u2014': '--',  # em dash
        '\u2026': '...',  # puntos suspensivos
        '\u00a0': ' ',  # espacio no separable
    }

    for old, new in replacements.items():
        texto = texto.replace(old, new)

    # Para idiomas RTL (árabe, hebreo), mantener caracteres Unicode
    if any(ord(char) > 1424 and ord(char) < 1791 for char in texto):  # Rango árabe/hebreo
        return texto

    # Eliminar caracteres no ASCII solo si no es texto RTL
    texto = re.sub(r'[^\x00-\x7F\u0590-\u05FF\u0600-\u06FF\u4e00-\u9fff]+', '?', texto)

    return texto


def extraer_texto_pdf(pdf_file):
    """
    Extrae texto de un archivo PDF usando PyMuPDF
    """
    try:
        # Leer el PDF
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        texto_completo = ""

        # Extraer texto de cada página
        for page_num in range(pdf_document.page_count):
            page = pdf_document[page_num]
            texto_pagina = page.get_text()
            texto_completo += f"\n--- Página {page_num + 1} ---\n"
            texto_completo += texto_pagina

        pdf_document.close()
        return texto_completo

    except Exception as e:
        st.error(f"Error al extraer texto del PDF: {str(e)}")
        return ""


def preprocesar_imagen(image):
    """
    Preprocesa la imagen para mejorar el OCR, especialmente para texto hebreo/árabe
    """
    try:
        # Convertir PIL a numpy array para OpenCV
        img_array = np.array(image)

        # Convertir a escala de grises si es necesario
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array

        # Aplicar filtro de desenfoque para suavizar
        blurred = cv2.GaussianBlur(gray, (1, 1), 0)

        # Binarización adaptativa para mejorar contraste
        binary = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)

        # Operaciones morfológicas para limpiar el texto
        kernel = np.ones((1, 1), np.uint8)
        cleaned = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

        # Convertir de vuelta a PIL Image
        processed_image = Image.fromarray(cleaned)

        return processed_image

    except Exception as e:
        st.warning(f"Error en preprocesamiento: {str(e)}")
        return image


def mejorar_imagen_para_ocr(image):
    """
    Mejora la imagen usando PIL para mejor OCR
    """
    try:
        # Redimensionar si es muy pequeña
        width, height = image.size
        if width < 300 or height < 300:
            factor = max(300 / width, 300 / height)
            new_size = (int(width * factor), int(height * factor))
            image = image.resize(new_size, Image.LANCZOS)

        # Convertir a escala de grises
        if image.mode != 'L':
            image = image.convert('L')

        # Mejorar contraste
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)

        # Mejorar nitidez
        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(1.5)

        # Aplicar umbralización
        threshold = 128
        image = image.point(lambda x: 255 if x > threshold else 0, mode='1')

        return image

    except Exception as e:
        st.warning(f"Error en mejora de imagen: {str(e)}")
        return image


def ocr_con_idioma_especifico(image, codigo_idioma, usar_preprocesamiento=True):
    """
    Realiza OCR con un idioma específico y preprocesamiento de imagen
    """
    try:
        # Crear copia de la imagen original
        imagen_procesada = image.copy()

        # Aplicar preprocesamiento si está habilitado
        if usar_preprocesamiento:
            imagen_mejorada = mejorar_imagen_para_ocr(imagen_procesada)
            imagen_opencv = preprocesar_imagen(imagen_mejorada)
        else:
            imagen_opencv = imagen_procesada

        # Configuración específica para diferentes tipos de idiomas
        if codigo_idioma in ['ara', 'heb']:  # Idiomas RTL
            # Configuraciones múltiples para hebreo/árabe
            configs = [
                r'--oem 3 --psm 6 -c preserve_interword_spaces=1',
                r'--oem 3 --psm 4 -c preserve_interword_spaces=1',
                r'--oem 3 --psm 3 -c preserve_interword_spaces=1',
                r'--oem 1 --psm 6 -c preserve_interword_spaces=1'
            ]
        elif codigo_idioma in ['chi_sim', 'chi_tra', 'jpn', 'kor']:  # Idiomas CJK
            configs = [
                r'--oem 3 --psm 6 -c textord_force_make_prop_words=F',
                r'--oem 3 --psm 4 -c textord_force_make_prop_words=F'
            ]
        else:  # Idiomas LTR normales
            configs = [r'--oem 3 --psm 6', r'--oem 3 --psm 4']

        if codigo_idioma == 'auto':
            # Intentar con múltiples idiomas
            idiomas_comunes = ['heb', 'ara', 'eng', 'chi_sim', 'spa', 'fra']
            mejor_resultado = ""

            for idioma in idiomas_comunes:
                try:
                    for config in configs[:2]:  # Solo las 2 primeras configuraciones
                        resultado = pytesseract.image_to_string(imagen_opencv, lang=idioma, config=config)
                        if len(resultado.strip()) > len(mejor_resultado.strip()):
                            mejor_resultado = resultado
                            if len(resultado.strip()) > 50:  # Si encontramos suficiente texto, parar
                                break
                except:
                    continue
                if len(mejor_resultado.strip()) > 50:
                    break

            return mejor_resultado
        else:
            # Probar diferentes configuraciones y quedarse con la mejor
            mejor_resultado = ""

            for config in configs:
                try:
                    resultado = pytesseract.image_to_string(imagen_opencv, lang=codigo_idioma, config=config)
                    if len(resultado.strip()) > len(mejor_resultado.strip()):
                        mejor_resultado = resultado
                except Exception as config_error:
                    st.warning(f"Error con configuración {config}: {str(config_error)}")
                    continue

            # Si no se obtuvo resultado, intentar con la imagen original
            if not mejor_resultado.strip():
                try:
                    mejor_resultado = pytesseract.image_to_string(image, lang=codigo_idioma, config=configs[0])
                except:
                    pass

            return mejor_resultado

    except Exception as e:
        st.error(f"Error en OCR con idioma {codigo_idioma}: {str(e)}")
        # Fallback a inglés con imagen original
        try:
            return pytesseract.image_to_string(image, lang='eng')
        except:
            return ""


def traducir_con_gpt(texto_original, idioma_origen="auto"):
    """
    Traduce texto usando GPT-4 con detección mejorada de idioma
    """
    llm = ChatOpenAI(model="gpt-4o-mini", api_key=OPENAI_API_KEY)

    system_message = """Eres un traductor profesional experto en múltiples idiomas incluyendo árabe, hebreo, chino, japonés y coreano. 
    Traduce el texto manteniendo el formato, contexto y significado original. 
    Si el texto contiene caracteres especiales o es de un idioma de escritura de derecha a izquierda, maneja la traducción apropiadamente."""

    question = f"""
                Detecta el idioma del siguiente texto y tradúcelo al inglés de manera precisa y natural:

                {texto_original}

                Proporciona la traducción directamente sin explicaciones adicionales.
                """
    try:
        prompt = ChatPromptTemplate.from_messages(
            [
                ("system", system_message),
                (
                    "human",
                    [
                        {"type": "text", "text": "{input}"},
                    ],
                ),
            ]
        )

        chain = prompt | llm
        response = chain.invoke({"input": question})

        st.write("**Traducción completada:**")
        st.write(response.content)
        st.write("------------------------------------")

        return response.content
    except Exception as e:
        st.error(f"Error al traducir texto: {str(e)}")
        return "Error al procesar el texto. Por favor, verifica tu API key de OpenAI y la conexión a internet."


def crear_pdf_con_reportlab(texto_original, texto_traducido):
    """
    Crea un PDF usando ReportLab que maneja mejor Unicode
    """
    buffer = io.BytesIO()

    # Crear documento PDF
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    try:
        # Limpiar textos
        texto_original_limpio = limpiar_texto_para_pdf(texto_original)
        texto_traducido_limpio = limpiar_texto_para_pdf(texto_traducido)

        # Título
        title = Paragraph("Texto Extraído y Traducido", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Texto original
        subtitle1 = Paragraph("Texto Original:", styles['Heading2'])
        story.append(subtitle1)
        story.append(Spacer(1, 6))

        original_para = Paragraph(texto_original_limpio.replace('\n', '<br/>'), styles['Normal'])
        story.append(original_para)
        story.append(Spacer(1, 12))

        # Texto traducido
        subtitle2 = Paragraph("Traducción al Inglés:", styles['Heading2'])
        story.append(subtitle2)
        story.append(Spacer(1, 6))

        translated_para = Paragraph(texto_traducido_limpio.replace('\n', '<br/>'), styles['Normal'])
        story.append(translated_para)

        # Construir PDF
        doc.build(story)

    except Exception as e:
        st.error(f"Error al crear PDF: {str(e)}")
        # Crear PDF simple como fallback
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = [Paragraph("Error al generar PDF con formato completo", styles['Normal'])]
        doc.build(story)

    return buffer


# Sidebar para configuración
with st.sidebar:
    st.header("⚙️ Configuración")

    # Selección de idioma para OCR
    idioma_seleccionado = st.selectbox(
        "Seleccionar idioma del texto:",
        options=list(IDIOMAS_TESSERACT.keys()),
        index=list(IDIOMAS_TESSERACT.keys()).index('Detección Automática')  # Hebreo por defecto para tu imagen
    )

    codigo_idioma = IDIOMAS_TESSERACT[idioma_seleccionado]

    # Opciones de preprocesamiento
    usar_preprocesamiento = st.checkbox(
        "🔧 Usar preprocesamiento de imagen",
        value=True,
        help="Mejora el contraste y limpia la imagen para mejor OCR"
    )

    mostrar_imagen_procesada = st.checkbox(
        "👁️ Mostrar imagen procesada",
        value=False,
        help="Muestra cómo se ve la imagen después del preprocesamiento"
    )

    # Verificar idiomas instalados
    if st.button("🔍 Verificar idiomas instalados"):
        idiomas_disponibles = verificar_idiomas_instalados()
        st.write("**Idiomas instalados:**")
        st.write(idiomas_disponibles)

        # Verificar si el idioma seleccionado está disponible
        if codigo_idioma != 'auto' and codigo_idioma not in idiomas_disponibles:
            st.warning(f"⚠️ El idioma {idioma_seleccionado} ({codigo_idioma}) no está instalado")
            if st.button(f"📥 Descargar {idioma_seleccionado}"):
                if descargar_paquete_idioma(codigo_idioma):
                    st.rerun()

    # Información sobre instalación manual
    with st.expander("📖 Instalación manual de idiomas"):
        st.markdown("""
        **Para instalar idiomas manualmente:**

        1. Descargar archivos .traineddata desde:
           https://github.com/tesseract-ocr/tessdata

        2. Copiar a la carpeta tessdata:
           - Windows: `C:\\Program Files\\Tesseract-OCR\\tessdata`
           - Linux: `/usr/share/tesseract-ocr/tessdata`
           - Mac: `/opt/homebrew/share/tessdata`

        **Códigos de idioma importantes:**
        - Árabe: ara
        - Hebreo: heb  
        - Chino Simplificado: chi_sim
        - Chino Tradicional: chi_tra
        - Japonés: jpn
        - Coreano: kor
        """)

# Subir archivo
uploaded_file = st.file_uploader("📤 Sube tu archivo aquí", type=["png", "jpg", "jpeg", "pdf"])

# Procesamiento del archivo subido
if uploaded_file:
    file_type = uploaded_file.type
    extracted_text = ""

    if file_type == "application/pdf":
        st.subheader("📄 Procesando PDF...")

        # Mostrar información del archivo
        st.write(f"**Nombre del archivo:** {uploaded_file.name}")
        st.write(f"**Tamaño:** {uploaded_file.size} bytes")

        # Extraer texto del PDF
        with st.spinner("📖 Extrayendo texto del PDF..."):
            extracted_text = extraer_texto_pdf(uploaded_file)

    else:  # Imagen
        st.subheader("🖼️ Procesando Imagen...")
        image = Image.open(uploaded_file)
        st.image(image, caption="📷 Imagen original", use_column_width=True)

        # Mostrar imagen procesada si está habilitado
        if mostrar_imagen_procesada and usar_preprocesamiento:
            with st.spinner("🔧 Preprocesando imagen..."):
                imagen_mejorada = mejorar_imagen_para_ocr(image.copy())
                imagen_final = preprocesar_imagen(imagen_mejorada)
                st.image(imagen_final, caption="🔧 Imagen procesada para OCR", use_column_width=True)

        # Extracción OCR con idioma específico
        with st.spinner(f"🧠 Extrayendo texto con OCR ({idioma_seleccionado})..."):
            try:
                extracted_text = ocr_con_idioma_especifico(image, codigo_idioma, usar_preprocesamiento)
            except Exception as e:
                st.error(f"Error en OCR: {str(e)}")
                if codigo_idioma != 'eng':
                    st.info("Intentando con inglés como fallback...")
                    extracted_text = ocr_con_idioma_especifico(image, 'eng', usar_preprocesamiento)
                else:
                    extracted_text = ""

        # Mostrar información adicional del OCR
        if extracted_text.strip():
            st.success(f"✅ Texto extraído exitosamente ({len(extracted_text)} caracteres)")
        else:
            st.warning("⚠️ No se pudo extraer texto. Intenta con:")
            st.markdown("""
            - Cambiar el idioma seleccionado
            - Activar/desactivar el preprocesamiento
            - Verificar que el idioma esté instalado
            - Usar una imagen de mejor calidad
            """)

    # Mostrar texto extraído
    if extracted_text.strip():
        st.subheader("📄 Texto extraído:")
        st.text_area("Texto original:", extracted_text, height=200)

        # Botón de traducción
        if st.button("🔄 Traducir Texto", type="primary"):
            with st.spinner("🤖 Traduciendo con GPT-4..."):
                translated_text = traducir_con_gpt(extracted_text, idioma_seleccionado)

            if translated_text:
                st.subheader("🔤 Traducción al inglés:")
                st.text_area("Texto traducido:", translated_text, height=300)

                # Botones para exportar
                st.markdown("### 💾 Exportar texto:")
                col1, col2, col3 = st.columns(3)

                # Exportar a PDF con ReportLab
                with col1:
                    try:
                        pdf_buffer = crear_pdf_con_reportlab(extracted_text, translated_text)
                        st.download_button(
                            label="📄 Descargar PDF",
                            data=pdf_buffer.getvalue(),
                            file_name="traduccion_ocr.pdf",
                            mime="application/pdf"
                        )
                    except Exception as e:
                        st.error(f"Error al generar PDF: {str(e)}")

                # Exportar a Word
                with col2:
                    try:
                        doc = Document()
                        doc.add_heading("Texto extraído y traducido", level=1)
                        doc.add_heading("Texto original:", level=2)
                        doc.add_paragraph(extracted_text)
                        doc.add_heading("Traducción al inglés:", level=2)
                        doc.add_paragraph(translated_text)
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        st.download_button(
                            label="📝 Descargar Word",
                            data=doc_buffer.getvalue(),
                            file_name="traduccion_ocr.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"Error al generar Word: {str(e)}")

                # Exportar a CSV
                with col3:
                    try:
                        df = pd.DataFrame({
                            "Idioma_Origen": [idioma_seleccionado],
                            "Texto_Original": [extracted_text],
                            "Texto_Traducido": [translated_text]
                        })
                        csv_buffer = io.StringIO()
                        df.to_csv(csv_buffer, index=False, encoding='utf-8-sig')  # BOM para Excel
                        st.download_button(
                            label="📊 Descargar CSV",
                            data=csv_buffer.getvalue(),
                            file_name="traduccion_ocr.csv",
                            mime="text/csv"
                        )
                    except Exception as e:
                        st.error(f"Error al generar CSV: {str(e)}")
    else:
        st.warning(
            "⚠️ No se pudo extraer texto del archivo. Verifica que el archivo contenga texto legible o prueba con un idioma diferente.")

# Información adicional
st.markdown("---")
st.markdown("### 📋 Información:")
st.markdown(f"""
- **Formatos soportados**: PNG, JPG, JPEG, PDF
- **Idiomas soportados**: {len(IDIOMAS_TESSERACT) - 1} idiomas incluidos árabe, hebreo, chino
- **Idioma seleccionado**: {idioma_seleccionado}
- **Exportación**: PDF, Word, CSV con soporte Unicode
- **OCR**: Optimizado para idiomas RTL y CJK
""")

# Consejos para mejor OCR
with st.expander("💡 Consejos para mejor reconocimiento"):
    st.markdown("""
    **Para mejorar la precisión del OCR:**

    1. **Calidad de imagen**: Usa imágenes de alta resolución (300 DPI mínimo)
    2. **Contraste**: Asegúrate de que el texto tenga buen contraste con el fondo
    3. **Idioma correcto**: Selecciona el idioma específico del texto
    4. **Texto horizontal**: El texto debe estar bien alineado
    5. **Sin ruido**: Evita fondos complejos o texto borroso
    6. **Preprocesamiento**: Activa la opción para imágenes con marcas de agua

    **Para texto hebreo específicamente:**
    - Asegúrate de tener instalado el paquete 'heb'
    - Usa el preprocesamiento para imágenes con fondo complejo
    - El texto debe estar claramente visible y sin superposiciones

    **Idiomas especiales:**
    - **Árabe/Hebreo**: Se procesan de derecha a izquierda automáticamente
    - **Chino/Japonés**: Soporta tanto caracteres simplificados como tradicionales
    - **Múltiples idiomas**: Usa "Detección Automática" para texto mixto
    """)